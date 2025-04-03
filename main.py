import sys

import openpyxl
import pymysql
from PyQt6.QtCore import Qt, QDate
from PyQt6.QtGui import QBrush, QColor, QPen, QFont, QPainter, QPixmap, QGuiApplication, QIcon, QTextCharFormat, \
    QTextBlockFormat, QTextCursor, QTextDocument, QTextTableFormat, QPageLayout, QFontDatabase
from PyQt6.QtPrintSupport import QPrinter, QPrintDialog
from PyQt6.QtWidgets import QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, \
    QPushButton, QMessageBox, QTableWidget, QTableWidgetItem, QComboBox, QDialog, QSizePolicy, QStyle, QHeaderView, \
    QFileDialog, QFormLayout, QDateEdit, QMenu
import socket
from datetime import datetime
import random
import string
from PyQt6.QtWidgets import QGroupBox
import os

from openpyxl.workbook import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm
from docx import Document
from docx.shared import Inches

class LoginWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.failed_attempts = 0
        self.captcha_text = ""
        self.setup_ui()
        self.center_window()  # Теперь этот метод существует
        self.set_fonts()
        self.set_window_icon()  # Добавляем метод для установки иконки

    def set_window_icon(self):
        """Устанавливает иконку для окна"""
        # Создаем QIcon из файла (если файл существует)
        app_icon = QIcon("logo.ico")

        # Если файл не найден, можно использовать встроенную иконку как fallback
        if app_icon.isNull():
            app_icon = self.style().standardIcon(QStyle.StandardPixmap.SP_MessageBoxInformation)

        self.setWindowIcon(app_icon)

    def center_window(self):
        """Центрирует окно на экране"""
        frame_geometry = self.frameGeometry()
        screen = QGuiApplication.primaryScreen().availableGeometry().center()
        frame_geometry.moveCenter(screen)
        self.move(frame_geometry.topLeft())

    def set_fonts(self):
        """Устанавливает шрифт Comic Sans MS для всех элементов"""
        comic_sans = QFont("Comic Sans MS", 10)

        # Устанавливаем шрифт для всех виджетов
        for widget in [
            self.username_label, self.username_input,
            self.password_label, self.password_input,
            self.toggle_password_button, self.captcha_group,
            self.captcha_input, self.refresh_captcha_button,
            self.login_button
        ]:
            widget.setFont(comic_sans)

        # Для CAPTCHA делаем шрифт крупнее
        captcha_font = QFont("Comic Sans MS", 16)
        self.captcha_display.setFont(captcha_font)

    def setup_ui(self):
        self.setWindowTitle("Авторизация")
        self.setGeometry(100, 100, 450, 200)  # Немного увеличили окно для шрифта

        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.layout = QVBoxLayout()
        self.layout.setContentsMargins(25, 25, 25, 25)
        self.layout.setSpacing(15)  # Увеличили промежутки между элементами

        # Логин
        self.username_label = QLabel("Логин:")
        self.username_input = QLineEdit()
        self.username_input.setFixedHeight(35)
        self.layout.addWidget(self.username_label)
        self.layout.addWidget(self.username_input)

        # Пароль
        self.password_label = QLabel("Пароль:")
        self.password_layout = QHBoxLayout()
        self.password_input = QLineEdit()
        self.password_input.setFixedHeight(35)
        self.password_input.setEchoMode(QLineEdit.EchoMode.Password)

        self.toggle_password_button = QPushButton()
        self.toggle_password_button.setFixedSize(35, 35)
        self.toggle_password_button.setIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_FileDialogDetailedView))
        self.toggle_password_button.setCheckable(True)
        self.toggle_password_button.toggled.connect(self.toggle_password_visibility)

        self.password_layout.addWidget(self.password_input)
        self.password_layout.addWidget(self.toggle_password_button)
        self.layout.addWidget(self.password_label)
        self.layout.addLayout(self.password_layout)

        # CAPTCHA элементы
        self.captcha_group = QGroupBox("Введите CAPTCHA")
        self.captcha_group.setVisible(False)
        self.captcha_group.setStyleSheet("""
            QGroupBox {
                background-color: rgb(118, 227, 131);
                border: 2px solid rgb(73, 140, 81);
                border-radius: 8px;
                margin-top: 15px;
                padding-top: 20px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 15px;
                padding: 0 5px;
            }
        """)

        self.captcha_layout = QVBoxLayout()

        # Виджет CAPTCHA
        self.captcha_display = QLabel()
        self.captcha_display.setFixedSize(350, 120)
        self.captcha_display.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.captcha_display.setStyleSheet("""
            background-color: white; 
            border: 2px solid rgb(73, 140, 81);
            border-radius: 5px;
            padding: 10px;
        """)

        # Кнопка обновления CAPTCHA
        self.refresh_captcha_button = QPushButton("Обновить CAPTCHA")
        self.refresh_captcha_button.setFixedHeight(40)
        self.refresh_captcha_button.setStyleSheet("""
            QPushButton {
                background-color: rgb(73, 140, 81);
                color: white;
                border: 2px solid rgb(50, 100, 50);
                border-radius: 8px;
                padding: 8px;
            }
            QPushButton:hover {
                background-color: rgb(85, 160, 85);
            }
            QPushButton:pressed {
                background-color: rgb(60, 120, 60);
            }
        """)
        self.refresh_captcha_button.clicked.connect(self.generate_captcha)

        # Поле ввода CAPTCHA
        self.captcha_input = QLineEdit()
        self.captcha_input.setFixedHeight(35)
        self.captcha_input.setPlaceholderText("Введите текст с картинки")

        self.captcha_layout.addWidget(self.captcha_display)
        self.captcha_layout.addWidget(self.refresh_captcha_button)
        self.captcha_layout.addWidget(self.captcha_input)
        self.captcha_group.setLayout(self.captcha_layout)
        self.layout.addWidget(self.captcha_group)

        # Кнопка входа
        self.login_button = QPushButton("Войти")
        self.login_button.setFixedHeight(45)
        self.login_button.setStyleSheet("""
            QPushButton {
                background-color: rgb(73, 140, 81);
                color: white;
                border: 2px solid rgb(50, 100, 50);
                border-radius: 8px;
                padding: 10px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: rgb(85, 160, 85);
            }
            QPushButton:pressed {
                background-color: rgb(60, 120, 60);
            }
        """)
        self.login_button.clicked.connect(self.login)

        self.layout.addWidget(self.login_button)
        self.central_widget.setLayout(self.layout)

    def generate_captcha(self):
        """Генерирует CAPTCHA с шрифтом Comic Sans MS"""
        self.captcha_text = ''.join(random.choices(
            string.ascii_uppercase + string.digits, k=4))

        pixmap = QPixmap(350, 120)
        pixmap.fill(Qt.GlobalColor.white)

        painter = QPainter(pixmap)

        # Устанавливаем шрифт Comic Sans MS для CAPTCHA
        font = QFont("Comic Sans MS")
        font.setPointSize(28)
        font.setBold(True)
        painter.setFont(font)

        # Рисуем символы CAPTCHA
        positions = [
            (20 + i * 70 + random.randint(-10, 10),
             70 + random.randint(-20, 20))
            for i in range(4)]

        for i, (x, y) in enumerate(positions):
            # Разные цвета для символов
            painter.setPen(QColor(
                random.randint(50, 150),
                random.randint(50, 150),
                random.randint(50, 150)))
            painter.drawText(x, y, self.captcha_text[i])

            # Добавляем перечеркивание
            if random.random() > 0.6:
                painter.setPen(QPen(QColor(255, 0, 0), 2))
                painter.drawLine(x - 15, y - 25, x + 25, y + 25)

        # Добавляем шум
        noise_font = QFont("Comic Sans MS", 10)
        painter.setFont(noise_font)
        for _ in range(30):
            painter.setPen(QColor(
                random.randint(150, 220),
                random.randint(150, 220),
                random.randint(150, 220)))
            painter.drawText(
                random.randint(0, 300), random.randint(0, 100),
                random.choice(string.ascii_letters + string.digits))

        painter.end()
        self.captcha_display.setPixmap(pixmap)

    def toggle_password_visibility(self, checked):
        """Переключает видимость пароля"""
        if checked:
            self.password_input.setEchoMode(QLineEdit.EchoMode.Normal)
            self.toggle_password_button.setIcon(
                self.style().standardIcon(QStyle.StandardPixmap.SP_FileDialogDetailedView))
        else:
            self.password_input.setEchoMode(QLineEdit.EchoMode.Password)
            self.toggle_password_button.setIcon(
                self.style().standardIcon(QStyle.StandardPixmap.SP_FileDialogDetailedView))

    def get_client_ip(self):
        """Получает IP-адрес клиента"""
        try:
            hostname = socket.gethostname()
            return socket.gethostbyname(hostname)
        except:
            return "127.0.0.1"

    def log_auth_attempt(self, user_id, login, user_type, success):
        """Логирует попытку авторизации в БД"""
        try:
            connection = pymysql.connect(host='localhost', user='root', password='', database='laboratoriya20')
            cursor = connection.cursor()

            cursor.execute("""
                INSERT INTO auth_history 
                (user_id, login, user_type, login_time, ip_address, success)
                VALUES (%s, %s, %s, %s, %s, %s)
            """, (user_id, login, user_type, datetime.now(), self.get_client_ip(), success))

            connection.commit()
            connection.close()
        except Exception as e:
            print(f"Ошибка при логировании авторизации: {e}")

    def login(self):
        username = self.username_input.text()
        password = self.password_input.text()

        if not username or not password:
            QMessageBox.warning(self, "Ошибка", "Заполните все поля")
            return

        # Если требуется CAPTCHA
        if self.failed_attempts > 0 and self.captcha_group.isVisible():
            if self.captcha_input.text().upper() != self.captcha_text:
                QMessageBox.warning(self, "Ошибка", "Неверная CAPTCHA")
                self.generate_captcha()
                return

        connection = pymysql.connect(host='localhost', user='root', password='', database='laboratoriya20')
        cursor = connection.cursor()

        # Проверка в таблице users
        cursor.execute("SELECT id, type FROM users WHERE login=%s AND password=%s", (username, password))
        user = cursor.fetchone()

        if user:
            user_id, user_type = user
            self.log_auth_attempt(user_id, username,
                                  'admin' if user_type == 1 else 'lab' if user_type == 2 else 'buh',
                                  True)

            if user_type == 2:
                self.lab_window = LabWindow(user_id)
                self.lab_window.show()
                self.close()
            elif user_type == 1:
                self.admin_window = AdminWindow(user_id)
                self.admin_window.show()
                self.close()
            elif user_type == 3:
                self.client_window = BuhWindow(user_id)
                self.client_window.show()
                self.close()
        else:
            # Проверка в таблице patients
            cursor.execute("SELECT id FROM patients WHERE login=%s AND pwd=%s", (username, password))
            patient = cursor.fetchone()

            if patient:
                patient_id = patient[0]
                self.log_auth_attempt(patient_id, username, 'patient', True)
                self.patient_window = PatientWindow(patient_id)
                self.patient_window.show()
                self.close()
            else:
                self.failed_attempts += 1
                self.log_auth_attempt(0, username, 'unknown', False)

                if self.failed_attempts >= 1:
                    # Показываем CAPTCHA после первой неудачной попытки
                    self.captcha_group.setVisible(True)
                    self.generate_captcha()
                    self.adjustSize()

                QMessageBox.warning(self, "Ошибка", "Неверный логин или пароль")

            connection.close()

    def toggle_password_visibility(self, checked):
        """Переключает видимость пароля"""
        if checked:
            self.password_input.setEchoMode(QLineEdit.EchoMode.Normal)
            self.toggle_password_button.setIcon(
                self.style().standardIcon(QStyle.StandardPixmap.SP_FileDialogDetailedView))
        else:
            self.password_input.setEchoMode(QLineEdit.EchoMode.Password)
            self.toggle_password_button.setIcon(
                self.style().standardIcon(QStyle.StandardPixmap.SP_FileDialogDetailedView))


class PatientWindow(QMainWindow):
    """Окно для пациента"""

    def __init__(self, patient_id):
        super().__init__()
        self.patient_id = patient_id
        self.setWindowTitle("Окно пациента")
        self.setGeometry(100, 30, 700, 700)
        self.setWindowIcon(QIcon('logo.ico'))

        # Установка цветов
        self.primary_bg = QColor(255, 255, 255)  # Основной фон (белый)
        self.secondary_bg = QColor(118, 227, 131)  # Дополнительный фон
        self.accent_color = QColor(73, 140, 81)  # Акцентный цвет

        # Настройка основного виджета
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.central_widget.setStyleSheet(f"background-color: {self.primary_bg.name()};")

        self.layout = QVBoxLayout()
        self.central_widget.setLayout(self.layout)

        # Установка шрифта
        font = QFont("Comic Sans MS", 10)
        QApplication.setFont(font)

        # Центрируем кнопку
        self.layout.addStretch()

        # Кнопка просмотра результатов с акцентным цветом
        self.view_results_button = QPushButton("Просмотреть мои анализы")
        self.view_results_button.setFixedSize(200, 50)
        self.view_results_button.setStyleSheet(
            f"background-color: {self.accent_color.name()};"
            "color: white;"
            "border-radius: 10px;"
        )
        self.view_results_button.clicked.connect(self.view_my_services)
        self.layout.addWidget(self.view_results_button, alignment=Qt.AlignmentFlag.AlignCenter)


        self.inf = QPushButton("Помощь")
        self.inf.setFixedSize(130, 50)
        self.inf.setStyleSheet(
            f"background-color: grey;"
            "color: white;"
            "border-radius: 10px;"
        )
        self.inf.clicked.connect(self.view)
        self.layout.addWidget(self.inf, alignment=Qt.AlignmentFlag.AlignCenter)

        self.layout.addStretch()

        # Кнопка "Назад" с дополнительным цветом фона
        self.back_button = QPushButton("Назад")
        self.back_button.setFixedHeight(30)
        self.back_button.setStyleSheet(
            f"background-color: {self.secondary_bg.name()};"
            "border-radius: 5px;"
        )
        self.back_button.clicked.connect(self.go_back)
        self.layout.addWidget(self.back_button)

    def view_my_services(self):
        self.view_my_services_window = ViewMyServicesWindow(self.patient_id)
        self.view_my_services_window.show()

    def view(self):
        """Показать результаты анализов пациента"""
        # Здесь можно реализовать просмотр результатов
        msg = QMessageBox()
        msg.setWindowTitle("Обратная связь")
        msg.setWindowIcon(QIcon('logo.ico'))
        msg.setText("В случае некорректной работы программы или при наличии вопросов обращайтесь по номеру:\n\n"
                    "+7(909)956-22-44\nили на почту:\nalekseevnapolina77@mail.ru\n")
        msg.setStyleSheet(f"font-family: 'Comic Sans MS';")
        msg.exec()

    def go_back(self):
        """Вернуться к окну авторизации"""
        self.login_window = LoginWindow()
        self.login_window.show()
        self.close()

class BloodCollectionWindow(QDialog):
    def __init__(self, user_id):
        super().__init__()
        self.user_id = user_id
        self.setWindowTitle("Забор крови")
        self.setGeometry(200, 200, 400, 100)
        self.setWindowIcon(QIcon('logo.ico'))
        self.patient_ids = {}  # Словарь для хранения соответствия логин -> ID

        self.layout = QFormLayout()
        self.setLayout(self.layout)

        # Patient selection
        self.patient_combo = QComboBox()
        self.load_patients()
        self.layout.addRow("Пациент:", self.patient_combo)

        # Tube barcode field with suggestion
        self.barcode_field = QLineEdit()
        self.barcode_field.setPlaceholderText("Введите номер пробирки")
        self.suggest_next_barcode()
        self.barcode_field.returnPressed.connect(self.validate_barcode)
        self.layout.addRow("Номер пробирки:", self.barcode_field)

        # Date field (default to today)
        self.date_field = QDateEdit()
        self.date_field.setDate(QDate.currentDate())
        self.date_field.setCalendarPopup(True)
        self.layout.addRow("Дата забора:", self.date_field)

        # Submit button
        self.submit_button = QPushButton("Подтвердить и создать штрих-код")
        self.submit_button.clicked.connect(self.submit_blood_collection)
        self.layout.addRow(self.submit_button)

    def load_patients(self):
        try:
            conn = pymysql.connect(host='localhost', user='root', password='', database='laboratoriya20')
            cursor = conn.cursor()
            cursor.execute("SELECT id, login FROM patients")
            patients = cursor.fetchall()

            # Сохраняем соответствие между логином и ID
            self.patient_ids = {}
            for patient_id, login in patients:
                self.patient_combo.addItem(login)
                self.patient_ids[login] = patient_id

            conn.close()
        except pymysql.Error as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить пациентов: {str(e)}")

    def suggest_next_barcode(self):
        try:
            conn = pymysql.connect(host='localhost', user='root', password='', database='laboratoriya20')
            cursor = conn.cursor()
            cursor.execute("SELECT MAX(barcode) FROM blood")
            last_barcode = cursor.fetchone()[0]
            next_barcode = str(int(last_barcode) + 1) if last_barcode else "100000"
            self.barcode_field.setText(next_barcode)
            conn.close()
        except pymysql.Error as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось получить следующий номер пробирки: {str(e)}")

    def validate_barcode(self):
        barcode = self.barcode_field.text()
        if not barcode.isdigit():
            QMessageBox.warning(self, "Ошибка", "Номер пробирки должен содержать только цифры")
            return False

        try:
            conn = pymysql.connect(host='localhost', user='root', password='', database='laboratoriya20')
            cursor = conn.cursor()
            cursor.execute("SELECT COUNT(*) FROM blood WHERE barcode = %s", (barcode,))
            count = cursor.fetchone()[0]
            conn.close()

            if count > 0:
                QMessageBox.warning(self, "Ошибка", "Этот номер пробирки уже используется")
                return False
            return True
        except pymysql.Error as e:
            QMessageBox.warning(self, "Ошибка", f"Ошибка проверки номера пробирки: {str(e)}")
            return False

    def submit_blood_collection(self):
        if not self.validate_barcode():
            return

        patient_login = self.patient_combo.currentText()
        barcode = self.barcode_field.text()
        date = self.date_field.date().toString("yyyy-MM-dd")

        try:
            conn = pymysql.connect(host='localhost', user='root', password='', database='laboratoriya20')
            cursor = conn.cursor()

            # Сначала получаем ID пациента по его логину
            cursor.execute("SELECT id FROM patients WHERE login = %s", (patient_login,))
            patient_id = cursor.fetchone()

            if not patient_id:
                QMessageBox.warning(self, "Ошибка", "Пациент не найден в базе данных")
                conn.close()
                return

            patient_id = patient_id[0]  # Извлекаем ID из кортежа

            # Теперь вставляем запись с ID пациента
            cursor.execute(
                "INSERT INTO blood (patient, barcode, date, lab_assistant_id) VALUES (%s, %s, %s, %s)",
                (patient_id, barcode, date, self.user_id)
            )
            conn.commit()
            conn.close()

            # Generate barcode
            self.generate_barcode_pdf(barcode, date)

            QMessageBox.information(self, "Успех", "Забор крови зарегистрирован. Штрих-код создан.")
            self.close()
        except pymysql.Error as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось сохранить данные: {str(e)}")

    def generate_barcode_pdf(self, barcode, date):
        # Generate unique code
        unique_code = ''.join([str(random.randint(0, 9)) for _ in range(6)])
        full_code = f"{barcode}{date.replace('-', '')}{unique_code}"

        # Create PDF directory
        os.makedirs('barcodes', exist_ok=True)
        filename = f"barcodes/barcode_{barcode}.pdf"

        # Calculate dimensions (in points)
        page_width = 80 * mm  # ~226.7 points
        page_height = 30 * mm  # ~85 points

        # Margins and spacing
        left_margin = 3.63 * mm
        right_margin = 2.31 * mm
        top_margin = 2 * mm
        spacing = 0.4 * mm

        # Bar dimensions
        bar_height = page_height - top_margin - 5 * mm  # Оставляем место для цифр
        short_bar_reduction = 8 * mm  # На сколько укорачиваем обычные штрихи

        c = canvas.Canvas(filename, pagesize=(page_width, page_height))

        # Calculate total barcode width needed
        unit_width = 0.1 * mm
        total_code_width = 0
        for digit in full_code:
            if digit == '0':
                total_code_width += 1.5 * mm
            else:
                total_code_width += (unit_width * int(digit)) + spacing

        # Scale factor to fit barcode in available width
        available_width = page_width - left_margin - right_margin
        scale_factor = available_width / total_code_width

        x = left_margin
        y = top_margin

        # Draw barcode
        for i, digit in enumerate(full_code):
            digit = int(digit)

            # Determine if this is a boundary bar
            is_boundary = (i < 2) or (i >= len(full_code) - 2) or \
                          (i == len(full_code) // 2 - 1) or (i == len(full_code) // 2)

            if digit == 0:
                # White space for zero
                x += 1.5 * mm * scale_factor
            else:
                # Calculate bar width
                bar_width = unit_width * digit * scale_factor

                # Draw bar
                if is_boundary:
                    # Full height boundary bars
                    c.rect(x, y, bar_width, bar_height, fill=1)
                else:
                    # Two-part bars: top (uniform) and bottom (variable)
                    uniform_part = bar_height - short_bar_reduction
                    c.rect(x, y, bar_width, uniform_part, fill=1)  # Top uniform part
                    c.rect(x, y + uniform_part, bar_width, short_bar_reduction, fill=1)  # Bottom variable part

                x += bar_width

            # Add spacing (except after last digit)
            if i < len(full_code) - 1:
                x += spacing * scale_factor

        # Draw human-readable numbers (centered below barcode)
        c.setFont("Helvetica-Bold", 6)
        text_width = c.stringWidth(full_code, "Helvetica-Bold", 6)
        text_x = left_margin + (available_width - text_width) / 2
        text_y = top_margin / 2  # Position below barcode with margin
        c.drawString(text_x, text_y, full_code)

        c.save()


class LabWindow(QMainWindow):
    def __init__(self, user_id):
        super().__init__()
        self.user_id = user_id
        self.setWindowTitle("Окно лаборанта")
        self.setGeometry(100, 30, 700, 800)  # Увеличили высоту для новой кнопки
        self.setWindowIcon(QIcon('logo.ico'))

        # Установка цветов
        self.primary_bg = QColor(255, 255, 255)
        self.secondary_bg = QColor(118, 227, 131)
        self.accent_color = QColor(73, 140, 81)

        # Настройка основного виджета
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.central_widget.setStyleSheet(f"background-color: {self.primary_bg.name()};")

        self.layout = QVBoxLayout()
        self.central_widget.setLayout(self.layout)

        # Установка шрифта
        font = QFont("Comic Sans MS", 10)
        QApplication.setFont(font)

        # Заголовок
        self.l = QLabel("\nДобро пожаловать!\nЛаборатория №20 рада Вас видеть.\n\n")
        self.l.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.l.setStyleSheet(f"color: {self.accent_color.name()}; font-weight: bold;")
        self.layout.addWidget(self.l)

        # Кнопка "Добавить услугу"
        self.add_service_button = QPushButton("Добавить услугу")
        self.add_service_button.setStyleSheet(
            f"background-color: {self.accent_color.name()};"
            "color: white;"
            "border-radius: 10px;"
            "padding: 18px;"
        )
        self.add_service_button.clicked.connect(self.open_add_service_window)
        self.layout.addWidget(self.add_service_button)

        # Кнопка "Просмотреть все услуги"
        self.view_services_button = QPushButton("Просмотреть все услуги")
        self.view_services_button.setStyleSheet(
            f"background-color: {self.accent_color.name()};"
            "color: white;"
            "border-radius: 10px;"
            "padding: 18px;"
        )
        self.view_services_button.clicked.connect(self.view_services)
        self.layout.addWidget(self.view_services_button)

        # Кнопка "Забор крови"
        self.blood_collection_button = QPushButton("Забор крови")
        self.blood_collection_button.setStyleSheet(
            f"background-color: {self.accent_color.name()};"
            "color: white;"
            "border-radius: 10px;"
            "padding: 18px;"
        )
        self.blood_collection_button.clicked.connect(self.open_blood_collection_window)
        self.layout.addWidget(self.blood_collection_button)

        self.change_status_button = QPushButton("Изменить статус услуги")
        self.change_status_button.setStyleSheet(
            f"background-color: {self.accent_color.name()};"
            "color: white;"
            "border-radius: 10px;"
            "padding: 18px;"
        )
        self.change_status_button.clicked.connect(self.open_change_status_window)
        self.layout.addWidget(self.change_status_button)

        self.change_zabor_button = QPushButton("Изменить статус забора биоматериала")
        self.change_zabor_button.setStyleSheet(
            f"background-color: {self.accent_color.name()};"
            "color: white;"
            "border-radius: 10px;"
            "padding: 18px;"
        )
        self.change_zabor_button.clicked.connect(self.open_change_zabor_window)
        self.layout.addWidget(self.change_zabor_button)

        self.change_res_button = QPushButton("Внести результат анализа")
        self.change_res_button.setStyleSheet(
            f"background-color: {self.secondary_bg.name()};"
            "color: white;"
            "border-radius: 10px;"
            "padding: 18px;"
        )
        self.change_res_button.clicked.connect(self.open_change_res_window)
        self.layout.addWidget(self.change_res_button)

        self.layout.addStretch()

        # Кнопка "Назад"
        self.back_button = QPushButton("Назад")
        self.back_button.setFixedHeight(30)
        self.back_button.setStyleSheet(
            f"background-color: {self.secondary_bg.name()};"
            "border-radius: 5px;"
        )
        self.back_button.clicked.connect(self.go_back)
        self.layout.addWidget(self.back_button)


    def open_add_service_window(self):
        self.add_service_window = AddServiceWindow(self.user_id)
        self.add_service_window.show()

    def view_services(self):
        self.view_services_window = ViewServicesWindow()
        self.view_services_window.show()

    def open_blood_collection_window(self):
        self.blood_collection_window = BloodCollectionWindow(self.user_id)
        self.blood_collection_window.show()

    def open_change_status_window(self):
        self.change_status_window = ChangeStatusWindow(self.user_id)
        self.change_status_window.show()

    def open_change_zabor_window(self):
        self.change_status_window = ChangeZaborWindow(self.user_id)
        self.change_status_window.show()

    def open_change_res_window(self):
        self.change_status_window = ChangeResWindow(self.user_id)
        self.change_status_window.show()

    def go_back(self):
        self.login_window = LoginWindow()
        self.login_window.show()
        self.close()


class ChangeStatusWindow(QDialog):
    def __init__(self, user_id):
        super().__init__()
        self.user_id = user_id
        self.setWindowTitle("Изменение статуса услуги")
        self.setGeometry(200, 200, 450, 200)
        self.setWindowIcon(QIcon('logo.ico'))

        self.layout = QVBoxLayout()
        self.setLayout(self.layout)

        # Выбор услуги
        self.service_label = QLabel("Выберите услугу:")
        self.service_combo = QComboBox()
        self.layout.addWidget(self.service_label)
        self.layout.addWidget(self.service_combo)

        # Выбор нового статуса
        self.status_label = QLabel("Новый статус:")
        self.status_combo = QComboBox()
        self.layout.addWidget(self.status_label)
        self.layout.addWidget(self.status_combo)

        # Кнопка подтверждения
        self.confirm_button = QPushButton("Изменить статус")
        self.confirm_button.clicked.connect(self.change_status)
        self.layout.addWidget(self.confirm_button)

        # Загрузка данных
        self.load_services()
        self.load_statuses()

    def load_services(self):
        try:
            conn = pymysql.connect(host='localhost', user='root', password='', database='laboratoriya20')
            cursor = conn.cursor()

            # Загружаем услуги из blood_services с информацией о биоматериале
            cursor.execute("""
                SELECT bs.id, s.name, b.barcode 
                FROM blood_services bs
                JOIN service s ON bs.service = s.id
                JOIN blood b ON bs.blood = b.id
            """)

            services = cursor.fetchall()
            for service in services:
                self.service_combo.addItem(f"{service[1]} (Пробирка: {service[2]})", service[0])

            conn.close()
        except pymysql.Error as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить услуги: {str(e)}")

    def load_statuses(self):
        try:
            conn = pymysql.connect(host='localhost', user='root', password='', database='laboratoriya20')
            cursor = conn.cursor()

            cursor.execute("SELECT id, name FROM status")
            statuses = cursor.fetchall()
            for status in statuses:
                self.status_combo.addItem(status[1], status[0])

            conn.close()
        except pymysql.Error as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить статусы: {str(e)}")

    def change_status(self):
        service_id = self.service_combo.currentData()
        new_status_id = self.status_combo.currentData()

        if not service_id or not new_status_id:
            QMessageBox.warning(self, "Ошибка", "Выберите услугу и новый статус")
            return

        try:
            conn = pymysql.connect(host='localhost', user='root', password='', database='laboratoriya20')
            cursor = conn.cursor()

            # Обновляем статус и устанавливаем finished=1 если статус "Завершено"
            cursor.execute("""
                UPDATE blood_services 
                SET status = %s, 
                    finished = CASE 
                        WHEN %s = (SELECT id FROM status WHERE name = 'Завершено') THEN 1 
                        ELSE finished 
                    END
                WHERE id = %s
            """, (new_status_id, new_status_id, service_id))

            conn.commit()
            QMessageBox.information(self, "Успех", "Статус услуги успешно изменен")
            self.close()
        except pymysql.Error as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось изменить статус: {str(e)}")
        finally:
            if conn:
                conn.close()


class ChangeZaborWindow(QDialog):
    def __init__(self, user_id):
        super().__init__()
        self.user_id = user_id
        self.setWindowTitle("Изменение статуса забора биоматериала")
        self.setGeometry(200, 200, 450, 200)
        self.setWindowIcon(QIcon('logo.ico'))

        self.layout = QVBoxLayout()
        self.setLayout(self.layout)

        # Выбор услуги
        self.service_label = QLabel("Выберите услугу:")
        self.service_combo = QComboBox()
        self.layout.addWidget(self.service_label)
        self.layout.addWidget(self.service_combo)

        # Выбор нового статуса
        self.status_label = QLabel("Принято:")
        self.status_combo = QComboBox()
        self.status_combo.addItem("да", 1)  # 1 соответствует True/да
        self.status_combo.addItem("нет", 0)  # 0 соответствует False/нет
        self.layout.addWidget(self.status_label)
        self.layout.addWidget(self.status_combo)

        # Кнопка подтверждения
        self.confirm_button = QPushButton("Изменить статус")
        self.confirm_button.clicked.connect(self.change_status)
        self.layout.addWidget(self.confirm_button)

        # Загрузка данных
        self.load_services()

    def load_services(self):
        try:
            conn = pymysql.connect(host='localhost', user='root', password='', database='laboratoriya20')
            cursor = conn.cursor()

            # Загружаем услуги из blood_services с информацией о биоматериале
            cursor.execute("""
                SELECT bs.id, s.name, b.barcode 
                FROM blood_services bs
                JOIN service s ON bs.service = s.id
                JOIN blood b ON bs.blood = b.id
            """)

            services = cursor.fetchall()
            for service in services:
                self.service_combo.addItem(f"{service[1]} (Пробирка: {service[2]})", service[0])

            conn.close()
        except pymysql.Error as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить услуги: {str(e)}")

    def change_status(self):
        if self.service_combo.currentIndex() == -1:
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите услугу")
            return

        try:
            service_id = self.service_combo.currentData()
            new_status = self.status_combo.currentData()  # Получаем 1 или 0

            conn = pymysql.connect(host='localhost', user='root', password='', database='laboratoriya20')
            cursor = conn.cursor()

            # Обновляем поле accepted в таблице blood_services
            cursor.execute("""
                UPDATE blood_services 
                SET accepted = %s 
                WHERE id = %s
            """, (new_status, service_id))

            conn.commit()
            conn.close()

            QMessageBox.information(self, "Успех", "Статус забора биоматериала успешно изменен")
            self.close()

        except pymysql.Error as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось изменить статус: {str(e)}")


class ChangeResWindow(QDialog):
    def __init__(self, user_id):
        super().__init__()
        self.user_id = user_id
        self.setWindowTitle("Изменение результата забора биоматериала")
        self.setGeometry(200, 200, 450, 200)
        self.setWindowIcon(QIcon('logo.ico'))

        self.layout = QVBoxLayout()
        self.setLayout(self.layout)

        # Выбор услуги
        self.service_label = QLabel("Выберите услугу:")
        self.service_combo = QComboBox()
        self.layout.addWidget(self.service_label)
        self.layout.addWidget(self.service_combo)

        # Поле для ввода результата
        self.result_label = QLabel("Введите результат:")
        self.result_input = QLineEdit()
        self.result_input.setPlaceholderText("Введите числовое значение...")
        self.layout.addWidget(self.result_label)
        self.layout.addWidget(self.result_input)

        # Кнопка подтверждения
        self.confirm_button = QPushButton("Изменить результат")
        self.confirm_button.clicked.connect(self.change_result)
        self.layout.addWidget(self.confirm_button)

        # Загрузка данных
        self.load_services()

    def load_services(self):
        try:
            conn = pymysql.connect(host='localhost', user='root', password='', database='laboratoriya20')
            cursor = conn.cursor()

            # Загружаем услуги из blood_services с информацией о биоматериале
            cursor.execute("""
                SELECT bs.id, s.name, b.barcode 
                FROM blood_services bs
                JOIN service s ON bs.service = s.id
                JOIN blood b ON bs.blood = b.id
            """)

            services = cursor.fetchall()
            for service in services:
                self.service_combo.addItem(f"{service[1]} (Пробирка: {service[2]})", service[0])

            conn.close()
        except pymysql.Error as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить услуги: {str(e)}")

    def change_result(self):
        if self.service_combo.currentIndex() == -1:
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите услугу")
            return

        result_text = self.result_input.text().strip()

        # Проверка на пустое значение
        if not result_text:
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, введите результат")
            return

        # Проверка на вещественное число
        try:
            result_value = float(result_text)
        except ValueError:
            QMessageBox.warning(self, "Ошибка", "Результат должен быть числом (целым или дробным)")
            return

        try:
            service_id = self.service_combo.currentData()

            conn = pymysql.connect(host='localhost', user='root', password='', database='laboratoriya20')
            cursor = conn.cursor()

            # Обновляем поле result в таблице blood_services
            cursor.execute("""
                UPDATE blood_services 
                SET result = %s 
                WHERE id = %s
            """, (result_value, service_id))

            conn.commit()
            conn.close()

            QMessageBox.information(self, "Успех", "Результат успешно изменен")
            self.close()

        except pymysql.Error as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось изменить результат: {str(e)}")

class AdminWindow(QMainWindow):
    def __init__(self, user_id):
        super().__init__()
        self.user_id = user_id
        self.setWindowTitle("Окно администратора")
        self.setGeometry(100, 100, 700, 700)
        self.setWindowIcon(QIcon('logo.ico'))

        # Цветовая схема
        self.primary_bg = QColor(255, 255, 255)    # Основной фон (белый)
        self.secondary_bg = QColor(118, 227, 131)  # Дополнительный фон
        self.accent_color = QColor(73, 140, 81)    # Акцентный цвет

        # Основной виджет
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.central_widget.setStyleSheet(f"background-color: {self.primary_bg.name()};")

        self.layout = QVBoxLayout()
        self.central_widget.setLayout(self.layout)

        # Установка шрифта
        font = QFont("Comic Sans MS", 10)
        QApplication.setFont(font)

        # Заголовок
        self.l = QLabel("Добро пожаловать!\nЛаборатория №20 рада Вас видеть.")
        self.l.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.l.setStyleSheet(f"""
            color: {self.accent_color.name()};
            font-weight: bold;
            font-size: 12px;
            margin-bottom: 20px;
        """)
        self.layout.addWidget(self.l)

        # Создаем стиль для основных кнопок
        button_style = f"""
            background-color: {self.accent_color.name()};
            color: white;
            border-radius: 8px;
            padding: 8px;
            min-height: 30px;
            margin: 5px;
        """

        # Кнопка "Просмотреть все услуги"
        self.view_services_button = QPushButton("Просмотреть все услуги")
        self.view_services_button.setStyleSheet(button_style)
        self.view_services_button.clicked.connect(self.view_services)
        self.layout.addWidget(self.view_services_button)

        # Кнопка "Добавить услугу"
        self.add_service_button = QPushButton("Добавить услугу")
        self.add_service_button.setStyleSheet(button_style)
        self.add_service_button.clicked.connect(self.open_add_service_window)
        self.layout.addWidget(self.add_service_button)

        # Кнопка "Просмотреть все виды услуг"
        self.view_all_services_button = QPushButton("Просмотреть все виды услуг")
        self.view_all_services_button.setStyleSheet(button_style)
        self.view_all_services_button.clicked.connect(self.view_all_services)
        self.layout.addWidget(self.view_all_services_button)

        # Кнопка "Добавить новый вид услуги"
        self.add_new_service_button = QPushButton("Добавить новый вид услуги")
        self.add_new_service_button.setStyleSheet(button_style)
        self.add_new_service_button.clicked.connect(self.open_add_new_service_window)
        self.layout.addWidget(self.add_new_service_button)

        # Кнопка "Просмотреть историю авторизаций"
        self.view_auth_history_button = QPushButton("Просмотреть историю авторизаций")
        self.view_auth_history_button.setStyleSheet(button_style)
        self.view_auth_history_button.clicked.connect(self.view_auth_history)
        self.layout.addWidget(self.view_auth_history_button)


        self.otch = QPushButton("Сформировать отчет")
        self.otch.setStyleSheet(button_style)
        self.otch.clicked.connect(self.otchet_month)
        self.layout.addWidget(self.otch)

        # Добавляем растягивающее пространство перед кнопкой "Назад"
        self.layout.addStretch()

        # Кнопка "Назад"
        self.back_button = QPushButton("Назад")
        self.back_button.setFixedHeight(50)
        self.back_button.setStyleSheet(f"""
            background-color: {self.secondary_bg.name()};
            border-radius: 5px;
            margin-top: 10px;
        """)
        self.back_button.clicked.connect(self.go_back)
        self.layout.addWidget(self.back_button)

    def otchet_month(self):
        self.month_window = OtchetMonth()
        self.month_window.show()

    def view_auth_history(self):
        """Открывает окно с историей авторизаций"""
        self.history_window = AuthHistoryWindow()
        self.history_window.show()

    def view_services(self):
        self.view_services_window = ViewServicesWindow()
        self.view_services_window.show()

    def open_add_service_window(self):
        self.add_service_window = AddServiceWindow(self.user_id)
        self.add_service_window.show()

    def open_add_new_service_window(self):
        self.add_new_service_window = AddNewServiceWindow()
        self.add_new_service_window.show()

    def view_all_services(self):
        self.view_all_services_window = ViewAllServicesWindow()
        self.view_all_services_window.show()

    def go_back(self):
        self.login_window = LoginWindow()
        self.login_window.show()
        self.close()

class OtchetMonth(QDialog):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Отчет за месяц")
        self.setGeometry(100, 100, 1100, 600)
        self.setWindowIcon(QIcon('logo.ico'))

        self.layout = QVBoxLayout()
        self.setLayout(self.layout)

        # Добавляем фильтры
        self.filter_layout = QHBoxLayout()

        # Поле для фильтра по услуге
        self.service_filter_label = QLabel("Фильтр по услуге:")
        self.service_filter_input = QLineEdit()
        self.service_filter_input.setPlaceholderText("Введите название услуги...")

        # Кнопка применения фильтра
        self.filter_button = QPushButton("Применить фильтр")
        self.filter_button.clicked.connect(self.load_data)

        self.filter_layout.addWidget(self.service_filter_label)
        self.filter_layout.addWidget(self.service_filter_input)
        self.filter_layout.addWidget(self.filter_button)
        self.layout.addLayout(self.filter_layout)

        # Создаем таблицу
        self.table = QTableWidget()
        self.table.setColumnCount(10)
        self.table.setHorizontalHeaderLabels([
            "ID", "Услуга", "Результат", "Дата завершения",
            "Принято", "Статус", "Анализатор", "Лаборант", "Пациент", "Пробирка"
        ])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        self.table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.layout.addWidget(self.table)

        # Кнопка экспорта с выпадающим меню
        self.export_button = QPushButton("Экспорт отчета")

        # Меню экспорта (только Excel и Word)
        self.export_menu = QMenu()
        self.export_menu.addAction("Сохранить как Excel", self.export_to_excel)
        self.export_menu.addAction("Сохранить как Word", self.export_to_word)

        # Устанавливаем меню на кнопку
        self.export_button.setMenu(self.export_menu)

        # Добавляем кнопку на макет
        self.layout.addWidget(self.export_button, alignment=Qt.AlignmentFlag.AlignRight)

        # Загружаем данные
        self.load_data()

    def load_data(self):
        try:
            connection = pymysql.connect(
                host='localhost',
                user='root',
                password='',
                database='laboratoriya20'
            )
            cursor = connection.cursor()

            service_filter = self.service_filter_input.text().strip()

            query = """
                SELECT bs.id, s.name, bs.result, bs.finished, bs.accepted, 
                       st.name, a.name, u.login, patient.login, b.barcode
                FROM blood_services bs
                JOIN service s ON bs.service = s.id
                JOIN status st ON bs.status = st.id
                JOIN analyzer a ON bs.analyzer = a.id
                JOIN users u ON bs.user = u.id
                JOIN blood b ON bs.blood = b.id
                JOIN patients patient ON b.patient = patient.id
                WHERE bs.finished >= DATE_SUB(NOW(), INTERVAL 30 DAY)
            """

            params = []
            if service_filter:
                query += " AND s.name LIKE %s"
                params.append(f"%{service_filter}%")

            query += " ORDER BY bs.finished DESC"

            cursor.execute(query, params)
            services = cursor.fetchall()

            self.table.setRowCount(len(services))
            for row, service in enumerate(services):
                for col, data in enumerate(service):
                    if col == 3 and data:
                        data = data.strftime("%Y-%m-%d %H:%M")
                    elif col == 4:
                        data = "Да" if data else "Нет"

                    item = QTableWidgetItem(str(data))
                    item.setFlags(item.flags() ^ Qt.ItemFlag.ItemIsEditable)
                    self.table.setItem(row, col, item)

            connection.close()

        except pymysql.Error as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить данные: {str(e)}")

    def export_to_excel(self):
        """Экспорт отчета в Excel файл"""

        file_path, _ = QFileDialog.getSaveFileName(
            self, "Сохранить отчет как Excel", "", "Excel Files (*.xlsx)"
        )

        if file_path:
            try:
                wb = Workbook()
                ws = wb.active
                ws.title = "Отчет за 30 дней"

                headers = ["ID", "Услуга", "Результат", "Дата завершения",
                           "Принято", "Статус", "Анализатор", "Лаборант",
                           "Пациент", "Пробирка"]
                ws.append(headers)

                for row in range(self.table.rowCount()):
                    row_data = []
                    for col in range(self.table.columnCount()):
                        item = self.table.item(row, col)
                        row_data.append(item.text() if item else "")
                    ws.append(row_data)

                wb.save(file_path)
                QMessageBox.information(self, "Успех", f"Excel успешно сохранен: {file_path}")

            except ImportError:
                QMessageBox.warning(self, "Ошибка",
                                    "Для экспорта в Excel установите openpyxl: pip install openpyxl")
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить Excel: {str(e)}")


    def export_to_word(self):
        """Экспорт отчета в Word файл"""
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Сохранить отчет как Word", "", "Word Files (*.docx)"
        )

        if file_path:
            try:
                doc = Document()
                doc.add_heading('Отчет за последние 30 дней', level=1)
                doc.add_paragraph(f'Дата формирования: {datetime.now().strftime("%Y-%m-%d %H:%M")}')

                table = doc.add_table(rows=1, cols=self.table.columnCount())
                hdr_cells = table.rows[0].cells

                for i in range(len(hdr_cells)):
                    hdr_cells[i].text = str(self.table.horizontalHeaderItem(i).text())

                for row in range(self.table.rowCount()):
                    row_cells = table.add_row().cells
                    for col in range(self.table.columnCount()):
                        item = self.table.item(row, col)
                        row_cells[col].text = item.text() if item else ""

                doc.save(file_path)
                QMessageBox.information(self, "Успех", f"Word успешно сохранен: {file_path}")

            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить Word: {str(e)}")

class AuthHistoryWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("История авторизаций")
        self.setGeometry(100, 100, 800, 600)
        self.setWindowIcon(QIcon('logo.ico'))

        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)

        self.layout = QVBoxLayout()

        # Панель фильтров
        self.filter_layout = QHBoxLayout()

        self.login_filter_label = QLabel("Фильтр по логину:")
        self.login_filter_input = QLineEdit()
        self.login_filter_input.setPlaceholderText("Введите логин...")

        self.date_sort_label = QLabel("Сортировка по дате:")
        self.date_sort_combo = QComboBox()
        self.date_sort_combo.addItems(["Новые сначала", "Старые сначала"])

        self.apply_filters_button = QPushButton("Применить")
        self.apply_filters_button.clicked.connect(self.load_history)

        self.filter_layout.addWidget(self.login_filter_label)
        self.filter_layout.addWidget(self.login_filter_input)
        self.filter_layout.addWidget(self.date_sort_label)
        self.filter_layout.addWidget(self.date_sort_combo)
        self.filter_layout.addWidget(self.apply_filters_button)
        self.export_button = QPushButton("Экспорт в CSV")
        self.export_button.clicked.connect(self.export_to_csv)
        self.filter_layout.addWidget(self.export_button)

        self.layout.addLayout(self.filter_layout)

        # Таблица для отображения истории
        self.table = QTableWidget()
        self.table.setColumnCount(5)
        self.table.setHorizontalHeaderLabels(["Время входа", "Логин", "Тип пользователя", "IP-адрес", "Статус"])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)

        self.layout.addWidget(self.table)
        self.central_widget.setLayout(self.layout)

        # Загружаем данные при открытии
        self.load_history()

    def load_history(self):
        """Загружает историю авторизаций с учетом фильтров"""
        try:
            connection = pymysql.connect(host='localhost', user='root', password='', database='laboratoriya20')
            cursor = connection.cursor()

            # Формируем SQL-запрос с учетом фильтров
            sql = "SELECT login_time, login, user_type, ip_address, success FROM auth_history"
            params = []

            # Фильтр по логину
            login_filter = self.login_filter_input.text().strip()
            if login_filter:
                sql += " WHERE login LIKE %s"
                params.append(f"%{login_filter}%")

            # Сортировка по дате
            if self.date_sort_combo.currentText() == "Новые сначала":
                sql += " ORDER BY login_time DESC"
            else:
                sql += " ORDER BY login_time ASC"

            cursor.execute(sql, params)
            history = cursor.fetchall()

            # Заполняем таблицу
            self.table.setRowCount(len(history))

            for row_idx, row in enumerate(history):
                login_time, login, user_type, ip_address, success = row

                # Форматируем время
                time_item = QTableWidgetItem(login_time.strftime("%Y-%m-%d %H:%M:%S"))

                # Форматируем статус
                status_item = QTableWidgetItem("Успешно" if success else "Ошибка")
                if not success:
                    status_item.setForeground(QBrush(QColor(255, 0, 0)))  # Красный цвет для ошибок

                self.table.setItem(row_idx, 0, time_item)
                self.table.setItem(row_idx, 1, QTableWidgetItem(login))
                self.table.setItem(row_idx, 2, QTableWidgetItem(user_type))
                self.table.setItem(row_idx, 3, QTableWidgetItem(ip_address))
                self.table.setItem(row_idx, 4, status_item)

            connection.close()
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить историю: {e}")

    def export_to_csv(self):
        """Экспортирует историю в CSV файл"""
        file_name, _ = QFileDialog.getSaveFileName(self, "Экспорт в CSV", "", "CSV Files (*.csv)")

        if file_name:
            try:
                with open(file_name, 'w', encoding='utf-8') as f:
                    f.write("Время входа;Логин;Тип пользователя;IP-адрес;Статус\n")

                    for row in range(self.table.rowCount()):
                        line = ';'.join([
                            self.table.item(row, 0).text(),
                            self.table.item(row, 1).text(),
                            self.table.item(row, 2).text(),
                            self.table.item(row, 3).text(),
                            self.table.item(row, 4).text()
                        ])
                        f.write(line + '\n')

                QMessageBox.information(self, "Успех", "Данные успешно экспортированы")
            except Exception as e:
                QMessageBox.warning(self, "Ошибка", f"Не удалось экспортировать данные: {e}")

class BuhWindow(QMainWindow):
    def __init__(self, user_id):
        super().__init__()
        self.user_id = user_id
        self.setWindowTitle("Окно бухгалтера")
        self.setGeometry(400, 30, 700, 700)
        self.setWindowIcon(QIcon('logo.ico'))
        self.primary_bg = QColor(255, 255, 255)    # Основной фон (белый)
        self.secondary_bg = QColor(118, 227, 131)  # Дополнительный фон
        self.accent_color = QColor(73, 140, 81)    # Акцентный цвет

        # Основной виджет
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.central_widget.setStyleSheet(f"background-color: {self.primary_bg.name()};")

        self.layout = QVBoxLayout()
        self.central_widget.setLayout(self.layout)

        # Установка шрифта
        font = QFont("Comic Sans MS", 10)
        QApplication.setFont(font)

        # Заголовок
        self.l = QLabel("Добро пожаловать!\nЛаборатория №20 рада Вас видеть.")
        self.l.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.l.setStyleSheet(f"""
            color: {self.accent_color.name()};
            font-weight: bold;
            font-size: 12px;
            margin-bottom: 20px;
        """)
        self.layout.addWidget(self.l)

        # Создаем стиль для основных кнопок
        button_style = f"""
            background-color: {self.accent_color.name()};
            color: white;
            border-radius: 8px;
            padding: 8px;
            min-height: 30px;
            margin: 5px;
        """

        self.view_auth_history_button = QPushButton("Просмотреть историю авторизаций")
        self.view_auth_history_button.setStyleSheet(button_style)
        self.view_auth_history_button.clicked.connect(self.view_auth_history)
        self.layout.addWidget(self.view_auth_history_button)

        self.otch = QPushButton("Сформировать отчет")
        self.otch.setStyleSheet(button_style)
        self.otch.clicked.connect(self.otchet_month)
        self.layout.addWidget(self.otch)


        self.back_button = QPushButton("Назад")
        self.back_button.setFixedHeight(30)
        self.back_button.clicked.connect(self.go_back)
        self.layout.addWidget(self.back_button)

        self.central_widget.setLayout(self.layout)

    def go_back(self):
        self.login_window = LoginWindow()
        self.login_window.show()
        self.close()

    def otchet_month(self):
        self.month_window = OtchetMonth()
        self.month_window.show()

    def view_auth_history(self):
        """Открывает окно с историей авторизаций"""
        self.history_window = AuthHistoryWindow()
        self.history_window.show()

class AddServiceWindow(QDialog):
    def __init__(self, user_id):
        super().__init__()
        self.user_id = user_id
        self.setWindowTitle("Добавить услугу")
        self.setGeometry(100, 100, 500, 500)  # Увеличиваем размер окна
        self.setWindowIcon(QIcon('logo.ico'))

        self.layout = QVBoxLayout()

        # Выбор биоматериала (пробирки)
        self.blood_label = QLabel("Биоматериал (пробирка):")
        self.blood_combo = QComboBox()
        self.layout.addWidget(self.blood_label)
        self.layout.addWidget(self.blood_combo)

        # Выбор услуги
        self.service_label = QLabel("Услуга:")
        self.service_combo = QComboBox()
        self.layout.addWidget(self.service_label)
        self.layout.addWidget(self.service_combo)

        # Поле результата
        self.result_label = QLabel("Результат:")
        self.result_input = QLineEdit()
        self.layout.addWidget(self.result_label)
        self.layout.addWidget(self.result_input)

        # Статус принятия
        self.accepted_label = QLabel("Принято (1 - да, 0 - нет):")
        self.accepted_input = QComboBox()  # Изменяем на выпадающий список
        self.accepted_input.addItem("Да", 1)
        self.accepted_input.addItem("Нет", 0)
        self.layout.addWidget(self.accepted_label)
        self.layout.addWidget(self.accepted_input)

        # Статус выполнения
        self.status_label = QLabel("Статус:")
        self.status_combo = QComboBox()
        self.layout.addWidget(self.status_label)
        self.layout.addWidget(self.status_combo)

        # Анализатор
        self.analyzer_label = QLabel("Анализатор:")
        self.analyzer_combo = QComboBox()
        self.layout.addWidget(self.analyzer_label)
        self.layout.addWidget(self.analyzer_combo)

        self.add_button = QPushButton("Добавить")
        self.add_button.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;  /* Зеленый цвет */
                color: white;
                border: none;
                border-radius: 8px;
                padding: 24px 24px;
                font-size: 14px;
                font-weight: bold;
                min-width: 100px;
            }

            QPushButton:hover {
                background-color: #45a049;  /* Темно-зеленый при наведении */
            }

            QPushButton:pressed {
                background-color: #3d8b40;  /* Еще темнее при нажатии */
                padding-top: 13px;
                padding-bottom: 11px;
            }

            QPushButton:disabled {
                background-color: #cccccc;  /* Серый при недоступности */
                color: #666666;
            }
        """)
        self.add_button.clicked.connect(self.add_service)
        self.layout.addWidget(self.add_button)

        self.back_button = QPushButton("Назад")
        self.back_button.clicked.connect(self.go_back)
        self.layout.addWidget(self.back_button)

        self.setLayout(self.layout)

        # Загрузка данных в выпадающие списки
        self.load_combos()

    def load_combos(self):
        try:
            connection = pymysql.connect(
                host='localhost',
                user='root',
                password='',
                database='laboratoriya20'
            )
            cursor = connection.cursor()

            # Загрузка доступных биоматериалов
            cursor.execute("SELECT id, barcode FROM blood")
            blood_samples = cursor.fetchall()
            for sample in blood_samples:
                self.blood_combo.addItem(f"Пробирка {sample[1]}", sample[0])

            # Загрузка услуг
            cursor.execute("SELECT id, name FROM service")
            services = cursor.fetchall()
            for service in services:
                self.service_combo.addItem(service[1], service[0])

            # Загрузка статусов
            cursor.execute("SELECT id, name FROM status")
            statuses = cursor.fetchall()
            for status in statuses:
                self.status_combo.addItem(status[1], status[0])

            # Загрузка анализаторов
            cursor.execute("SELECT id, name FROM analyzer")
            analyzers = cursor.fetchall()
            for analyzer in analyzers:
                self.analyzer_combo.addItem(analyzer[1], analyzer[0])

            connection.close()
        except pymysql.Error as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить данные: {str(e)}")

    def add_service(self):
        # Получаем данные из полей
        blood_id = self.blood_combo.currentData()
        service_id = self.service_combo.currentData()
        result = self.result_input.text()
        accepted = self.accepted_input.currentData()
        status_id = self.status_combo.currentData()
        analyzer_id = self.analyzer_combo.currentData()

        # Проверка заполнения полей
        if not all([blood_id, service_id, result, accepted is not None, status_id, analyzer_id]):
            QMessageBox.warning(self, "Ошибка", "Заполните все обязательные поля")
            return

        try:
            connection = pymysql.connect(
                host='localhost',
                user='root',
                password='',
                database='laboratoriya20'
            )
            cursor = connection.cursor()

            # Вставка данных в таблицу blood_services
            cursor.execute("""
                INSERT INTO blood_services 
                (service, result, accepted, status, analyzer, user, blood)
                VALUES (%s, %s, %s, %s, %s, %s, %s)
            """, (
                service_id,
                result,
                accepted,
                status_id,
                analyzer_id,
                self.user_id,  # Используем user_id из параметров
                blood_id
            ))

            connection.commit()
            QMessageBox.information(self, "Успех", "Услуга успешно добавлена")
            self.close()
        except pymysql.Error as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось добавить услугу: {str(e)}")
        finally:
            if connection:
                connection.close()

    def go_back(self):
        self.close()


class AddNewServiceWindow(QDialog):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Добавить новый вид услуги")
        self.setGeometry(100, 100, 350, 300)  # Увеличили высоту окна для новых полей
        self.setWindowIcon(QIcon('logo.ico'))

        self.layout = QVBoxLayout()

        # Название услуги
        self.name_label = QLabel("Название услуги:")
        self.name_input = QLineEdit()
        self.layout.addWidget(self.name_label)
        self.layout.addWidget(self.name_input)

        # Цена услуги
        self.price_label = QLabel("Цена услуги:")
        self.price_input = QLineEdit()
        self.price_input.setPlaceholderText("Введите число (например: 1500.50)")
        self.layout.addWidget(self.price_label)
        self.layout.addWidget(self.price_input)

        # Срок выполнения (в минутах)
        self.srok_label = QLabel("Срок выполнения (минут):")
        self.srok_input = QLineEdit()
        self.srok_input.setPlaceholderText("Целое число (например: 120)")
        self.layout.addWidget(self.srok_label)
        self.layout.addWidget(self.srok_input)

        # Норма отклонения
        self.otklon_label = QLabel("Норма отклонения:")
        self.otklon_input = QLineEdit()
        self.otklon_input.setPlaceholderText("Вещественное число (например: 0.5)")
        self.layout.addWidget(self.otklon_label)
        self.layout.addWidget(self.otklon_input)

        # Кнопка добавления
        self.add_button = QPushButton("Добавить")
        self.add_button.clicked.connect(self.add_new_service)
        self.layout.addWidget(self.add_button)

        # Кнопка назад
        self.back_button = QPushButton("Назад")
        self.back_button.setFixedHeight(30)
        self.back_button.clicked.connect(self.go_back)
        self.layout.addWidget(self.back_button)

        self.setLayout(self.layout)

    def add_new_service(self):
        # Получаем данные из полей ввода
        name = self.name_input.text().strip()
        price_text = self.price_input.text().strip()
        srok_text = self.srok_input.text().strip()
        otklon_text = self.otklon_input.text().strip()

        # Проверка заполнения всех полей
        if not all([name, price_text, srok_text, otklon_text]):
            QMessageBox.warning(self, "Ошибка", "Заполните все поля")
            return

        try:
            # Проверка и преобразование цены
            price = float(price_text)
            if price <= 0:
                raise ValueError("Цена должна быть положительной")

            # Проверка и преобразование срока выполнения
            srok = int(srok_text)
            if srok <= 0:
                raise ValueError("Срок выполнения должен быть положительным числом")

            # Проверка и преобразование нормы отклонения
            otklon = float(otklon_text)
            if otklon < 0:
                raise ValueError("Норма отклонения не может быть отрицательной")

        except ValueError as e:
            QMessageBox.warning(self, "Ошибка", f"Некорректные данные: {str(e)}")
            return

        # Подключение к базе данных и добавление записи
        try:
            connection = pymysql.connect(
                host='localhost',
                user='root',
                password='',
                database='laboratoriya20'
            )
            cursor = connection.cursor()

            cursor.execute(
                "INSERT INTO service (name, price, srok, otklon) VALUES (%s, %s, %s, %s)",
                (name, price, srok, otklon)
            )

            connection.commit()
            connection.close()

            QMessageBox.information(self, "Успех", "Новый вид услуги успешно добавлен")
            self.close()

        except pymysql.Error as e:
            QMessageBox.warning(self, "Ошибка базы данных", f"Не удалось добавить услугу: {str(e)}")

    def go_back(self):
        self.close()

class ViewServicesWindow(QDialog):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Просмотр всех услуг")
        self.setGeometry(100, 100, 800, 600)
        self.setWindowIcon(QIcon('logo.ico'))

        self.layout = QVBoxLayout()

        self.table = QTableWidget()
        self.layout.addWidget(self.table)

        self.back_button = QPushButton("Назад")
        self.back_button.setFixedHeight(30)
        self.back_button.clicked.connect(self.go_back)
        self.layout.addWidget(self.back_button)

        self.setLayout(self.layout)

        self.load_services()

    def load_services(self):
        connection = pymysql.connect(host='localhost', user='root', password='', database='laboratoriya20')
        cursor = connection.cursor()
        cursor.execute("""
            SELECT bs.id, s.name, bs.result, bs.accepted, st.name, a.name, u.login
            FROM blood_services bs
            JOIN service s ON bs.service = s.id
            JOIN status st ON bs.status = st.id
            JOIN analyzer a ON bs.analyzer = a.id
            JOIN users u ON bs.user = u.id
            ORDER BY bs.id
        """)
        services = cursor.fetchall()

        self.table.setRowCount(len(services))
        self.table.setColumnCount(7)
        self.table.setHorizontalHeaderLabels(["ID", "Услуга", "Результат", "Принято", "Статус", "Анализатор", "Лаборант"])

        for row, service in enumerate(services):
            for col, data in enumerate(service):
                self.table.setItem(row, col, QTableWidgetItem(str(data)))

        connection.close()

    def go_back(self):
        self.close()

class ViewAllServicesWindow(QDialog):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Просмотр всех видов услуг")
        self.setGeometry(100, 100, 555, 600)
        self.setWindowIcon(QIcon('logo.ico'))

        self.layout = QVBoxLayout()

        self.table = QTableWidget()
        self.layout.addWidget(self.table)

        self.back_button = QPushButton("Назад")
        self.back_button.setFixedHeight(30)
        self.back_button.clicked.connect(self.go_back)
        self.layout.addWidget(self.back_button)

        self.setLayout(self.layout)

        self.load_all_services()

    def load_all_services(self):
        connection = pymysql.connect(host='localhost', user='root', password='', database='laboratoriya20')
        cursor = connection.cursor()
        cursor.execute("SELECT id, name, price, srok, otklon FROM service")
        services = cursor.fetchall()

        self.table.setRowCount(len(services))
        self.table.setColumnCount(5)
        self.table.setHorizontalHeaderLabels(["ID", "Название", "Цена", "Срок выполнения(мин.)", "Норма"])

        for row, service in enumerate(services):
            for col, data in enumerate(service):
                self.table.setItem(row, col, QTableWidgetItem(str(data)))

        connection.close()

    def go_back(self):
        self.close()


class ViewMyServicesWindow(QDialog):
    def __init__(self, user_id):
        super().__init__()
        self.user_id = user_id
        self.setWindowTitle("Просмотр моих анализов")
        self.setGeometry(100, 100, 1040, 600)
        self.setWindowIcon(QIcon('logo.ico'))

        self.layout = QVBoxLayout()

        # Добавляем поля для фильтрации
        self.filter_layout = QHBoxLayout()

        # Фильтр по биоматериалу
        self.barcode_filter_label = QLabel("Биоматериал:")
        self.barcode_filter_input = QLineEdit()
        self.barcode_filter_input.setPlaceholderText("Номер биоматериала...")

        # Фильтр по типу услуги
        self.service_filter_label = QLabel("Услуга:")
        self.service_filter_input = QLineEdit()
        self.service_filter_input.setPlaceholderText("Тип услуги...")

        self.filter_button = QPushButton("Применить фильтры")
        self.filter_button.clicked.connect(self.load_my_services)

        self.filter_layout.addWidget(self.barcode_filter_label)
        self.filter_layout.addWidget(self.barcode_filter_input)
        self.filter_layout.addWidget(self.service_filter_label)
        self.filter_layout.addWidget(self.service_filter_input)
        self.filter_layout.addWidget(self.filter_button)
        self.layout.addLayout(self.filter_layout)

        self.table = QTableWidget()
        self.layout.addWidget(self.table)

        self.back_button = QPushButton("Назад")
        self.back_button.setFixedHeight(30)
        self.back_button.clicked.connect(self.go_back)
        self.layout.addWidget(self.back_button)

        self.setLayout(self.layout)

        self.load_my_services()

    def load_my_services(self):
        connection = pymysql.connect(host='localhost', user='root', password='', database='laboratoriya20')
        cursor = connection.cursor()

        # Получаем значения фильтров
        barcode_filter = self.barcode_filter_input.text().strip()
        service_filter = self.service_filter_input.text().strip()

        # Базовый запрос
        query = """
            SELECT bs.id, s.name, bs.result, bs.accepted, st.name, a.name, u.login, patient.login, b.barcode, bs.finished
            FROM blood_services bs
            JOIN service s ON bs.service = s.id
            JOIN status st ON bs.status = st.id
            JOIN analyzer a ON bs.analyzer = a.id
            JOIN users u ON bs.user = u.id
            JOIN blood b ON bs.blood = b.id
            JOIN patients patient ON b.patient = patient.id
            WHERE b.patient = %s
        """

        # Добавляем условия фильтрации
        params = [self.user_id]
        conditions = []

        if barcode_filter:
            conditions.append("b.barcode LIKE %s")
            params.append(f"%{barcode_filter}%")

        if service_filter:
            conditions.append("s.name LIKE %s")
            params.append(f"%{service_filter}%")

        if conditions:
            query += " AND " + " AND ".join(conditions)

        query += " ORDER BY b.barcode"

        cursor.execute(query, params)
        services = cursor.fetchall()

        self.table.setRowCount(len(services))
        self.table.setColumnCount(10)  # Увеличили количество столбцов до 10
        self.table.setHorizontalHeaderLabels(
            ["ID", "Услуга", "Результат", "Принято", "Статус", "Анализатор", "Лаборант", "Пациент", "Биоматериал",
             "Дата"])

        for row, service in enumerate(services):
            for col, data in enumerate(service):
                # Форматируем дату, если это столбец finished (индекс 9)
                if col == 9 and data:
                    data = data.strftime("%Y-%m-%d %H:%M") if data else ""
                self.table.setItem(row, col, QTableWidgetItem(str(data)))

        connection.close()

    def go_back(self):
        self.close()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    login_window = LoginWindow()
    login_window.show()
    sys.exit(app.exec())